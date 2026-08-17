import sys, os
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document

# Try Pintura Rev01
candidates = [
    r"C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\APRs - EDISER\PEX-JFX-ENG-027-PINTURA-Rev01.docx",
    r"C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\APRs - EDISER\PEX-JFX-ENG-034-TRANSPORTE_PALETEIRA_3TON.docx",
    r"C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\APRs - EDISER\PEX-JFX-ENG-028-FURO-LAJE-PERFURATRIZ-Rev_B.docx",
]

for p in candidates:
    if os.path.exists(p):
        print(f"\n\n=== {os.path.basename(p)} ===")
        doc = Document(p)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[P{i}] {para.style.name!r} | {para.text[:200]}")
        print("--- TABLES ---")
        for ti, table in enumerate(doc.tables):
            print(f"Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
            for ri, row in enumerate(table.rows):
                cells = [c.text.strip()[:80] for c in row.cells]
                print(f"  R{ri}: {cells}")
        break
    else:
        print(f"NOT FOUND: {p}")
