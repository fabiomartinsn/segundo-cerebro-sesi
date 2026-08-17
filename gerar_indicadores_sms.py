from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from jfx_docx_padrao import aplicar_rodape

NAVY   = RGBColor(0x0F, 0x2D, 0x52)
ORANGE = RGBColor(0xE0, 0x7B, 0x2A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x77, 0x77, 0x77)
FONT   = "Calibri"

def set_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6)
    tcPr.append(shd)

def ctext(cell, text, bold=False, color=None, size=9,
          italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color or RGBColor(0x1A, 0x1A, 0x1A)

def add_section_title(doc, text, color=None, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color or NAVY

def add_subsection(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.color.rgb = ORANGE

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY

def divider(doc, color="0F2D52"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"),   "single")
    bt.set(qn("w:sz"),    "6")
    bt.set(qn("w:space"), "1")
    bt.set(qn("w:color"), color)
    pBdr.append(bt)
    pPr.append(pBdr)

def make_table(doc, headers, rows_data, col_widths,
               header_bg="0F2D52", alt_bg=("F4F7FB", "FFFFFF")):
    """
    rows_data: list of rows; each row = list of (text, kwargs_dict).
    kwargs_dict keys: bold, color, size, italic, align.
    """
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = Cm(w)
    for i, h in enumerate(headers):
        ctext(tbl.cell(0, i), h, bold=True, color=WHITE, size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        set_bg(tbl.cell(0, i), header_bg)
    for ri, row_data in enumerate(rows_data):
        row = tbl.add_row()
        bg = alt_bg[ri % 2]
        for ci, (val, kw) in enumerate(row_data):
            kw2 = dict(kw)
            align = kw2.pop("align", WD_ALIGN_PARAGRAPH.LEFT)
            ctext(row.cells[ci], val, align=align, **kw2)
            set_bg(row.cells[ci], bg)
    return tbl

# ── Dados ──────────────────────────────────────────────────

IND_SEG_REATIVO = [
    ("TFCA",  "Taxa de Frequência de Acidentados Com Afastamento",
     "NCA × 10⁶ ÷ HHER",               "= 0", "Mensal",     "Menor", ""),
    ("TFSA",  "Taxa de Frequência de Acidentados Sem Afastamento",
     "NSA × 10⁶ ÷ HHER",               "= 0", "Mensal",     "Menor", "Também denominada 'Tussas'"),
    ("TOR",   "Taxa de Ocorrências Registráveis",
     "NAR × 10⁶ ÷ HHER",               "= 0", "Mensal",     "Menor", ""),
    ("TAR",   "Taxa de Acidentados Registráveis",
     "NAR (excl. PS) × 10⁶ ÷ HHER",    "= 0", "Mensal",     "Menor", "Exclui ocorrências de primeiros socorros (Classe 1)"),
    ("TG",    "Taxa de Gravidade",
     "TCA × 10⁶ ÷ HHER",               "= 0", "Trimestral", "Menor", ""),
    ("TAG",   "Taxa de Acidentados Graves",
     "NAR₄₅ × 10⁶ ÷ HHER",             "= 0", "Mensal",     "Menor", "Inclui apenas Classes 4 e 5: incap. permanente e morte"),
    ("TAF",   "Taxa de Acidentados Fatais",
     "NAF × 10⁸ ÷ HHER",               "= 0", "Mensal",     "Menor", "Fator 10⁸ — diferente dos demais indicadores"),
    ("NAF",   "Número de Acidentados Fatais",
     "Nº absoluto",                    "= 0", "Mensal",     "Menor", ""),
    ("TAP",   "Taxa de Anomalias com Alto Potencial",
     "NAP × 10⁶ ÷ HHER",               "= 0", "Mensal",     "Menor", ""),
    ("TFAT",  "Taxa de Frequência de Acidentes de Trânsito",
     "(NAL + NAPt) × 10⁶ ÷ (KML + KMP)", "= 0", "Mensal",  "Menor",
     "NAL/NAPt = acidentados em veículos leves/pesados; KML/KMP = km rodados"),
]

IND_SEG_PROATIVO = [
    ("IAB",      "Índice de Implementação de Ações de Abrangência",
     "Ações implantadas ÷ Ações estabelecidas × 100",         "—", "Trimestral", "Maior", ""),
    ("IARAR",    "Índice de Acompanhamento das Recomendações de Análise de Riscos",
     "(NREI + NRAPO + NRAPP + NRNJ) × 100 ÷ NTR",           "—", "Mensal",     "Maior", ""),
    ("PARC",     "Percentual de Análises de Riscos em Conformidade",
     "NARC × 100 ÷ NTAR",                                     "—", "Trimestral", "Maior", ""),
    ("TASP_T1",  "Taxa de Anomalias de Segurança de Processo — Nível 1",
     "NASP_T1 × 10⁶ ÷ HHERPR",                               "= 0", "Mensal",   "Menor", ""),
    ("TASP_T2",  "Taxa de Anomalias de Segurança de Processo — Nível 2",
     "NASP_T2 × 10⁶ ÷ HHERPR",                               "= 0", "Mensal",   "Menor", ""),
]

IND_CONTINGENCIA = [
    ("IRS",  "Índice de Realização de Simulados",
     "ΣSR_C ÷ ΣSP_C × 100",   "—", "Anual", "Maior",
     "SR_C = simulados realizados; SP_C = simulados planejados. Máximo: 100%."),
    ("IPAS", "Índice de Plano de Ação Pós-Simulado",
     "ΣAVR_C ÷ ΣAVP_C × 100", "—", "Anual", "Maior",
     "AVR = ações verificadas realizadas; AVP = ações verificadas planejadas. Máximo: 100%."),
]

IND_AMBIENTAL = [
    ("RSNPA",      "Resíduos Sólidos Não Perigosos Acumulados",
     "Acumulado desde início do contrato", "ton", "Mensal", "Menor", ""),
    ("RSNPG",      "Resíduos Sólidos Não Perigosos Gerados",
     "Gerado no mês de referência",        "ton", "Mensal", "Menor", ""),
    ("RSNPD_NRRR", "Resíduos Sólidos NP Destinados — rotas Não-RRR",
     "Destinado fora das rotas RRR",       "ton", "Mensal", "Menor", "Rotas preferidas: Reciclagem, Recuperação e Reúso"),
    ("RSPA",       "Resíduos Sólidos Perigosos Acumulados",
     "Acumulado desde início do contrato", "ton", "Mensal", "Menor", ""),
    ("RSPG",       "Resíduos Sólidos Perigosos Gerados",
     "Gerado no mês de referência",        "ton", "Mensal", "Menor", ""),
    ("RSPD_NRRR",  "Resíduos Sólidos Perigosos Destinados — rotas Não-RRR",
     "Destinado fora das rotas RRR",       "ton", "Mensal", "Menor", ""),
    ("ADC",        "Água Doce Captada",
     "Volume mensal captado",              "m³",  "Mensal", "Menor", ""),
    ("VAR",        "Volume de Água Reusada",
     "Volume mensal reusado",              "m³",  "Mensal", "Maior", ""),
    ("VAZO",       "Volume Vazado de Óleo e Derivados",
     "Volume acumulado no mês",            "m³",  "Mensal", "Menor",
     "Computa apenas volumes > 1 barril (0,159 m³) que atingiram corpo hídrico ou solo não impermeabilizado"),
]

IND_SAUDE = [
    ("ASO",    "Atestados de Saúde Ocupacional Válidos",
     "NASOV ÷ NEA × 100",                    "Maior", "Mensal",
     "Válido: exame clínico admissional ou periódico ≤ 365 dias"),
    ("PTP_S",  "Percentual de Tempo Perdido Saúde",
     "HS ÷ HP × 100",                        "Menor", "Mensal",
     "HS = horas perdidas por doença/saúde; HP = horas programadas"),
    ("TIDO",   "Taxa de Incidência de Doença Ocupacional",
     "Casos novos ÷ Efetivo ativo × 1.000",  "Menor", "Mensal", ""),
    ("EFA",    "Empregados Fisicamente Ativos",
     "Ativos (IPAQ) ÷ Efetivo ativo × 100",  "Maior", "Mensal",
     "Mensurado pelo questionário IPAQ"),
    ("IRREEA", "Índice de Resposta das Equipes de Evacuação Aeromédica",
     "Apresentações ≤ 20 min ÷ Acionamentos × 100", "Maior", "Mensal",
     "Meta ≥ 90%. Aplicável a unidades marítimas"),
]

REM_15 = [
    ("1",  "TFCA",                              "= 0",        "NCA × 10⁶ ÷ HHER"),
    ("2",  "TFSA",                              "= 0",        "NSA × 10⁶ ÷ HHER"),
    ("3",  "TG",                                "= 0",        "TCA × 10⁶ ÷ HHER"),
    ("4",  "TOR",                               "= 0",        "NAR × 10⁶ ÷ HHER"),
    ("5",  "TAR",                               "= 0",        "NAR (excl. PS) × 10⁶ ÷ HHER"),
    ("6",  "Acidentes com Morte",               "= 0",        "Nº absoluto"),
    ("7",  "Incidentes de Alto Potencial",      "= 0",        "Nº absoluto"),
    ("8",  "Desvios Críticos / Sistêmicos",     "< 2/mês",    "Nº absoluto"),
    ("9",  "Ocorrências Ambientais",            "= 0",        "Nº absoluto"),
    ("10", "HHER",                              "100% reg.",  "Valor real apurado (Roberta Muniz)"),
    ("11", "DDS Realizados / Programados",      "≥ 95%",      "Nº DDS ÷ dias úteis × 100"),
    ("12", "VCP Executadas / Programadas",      "= 100%",     "VCPs realizadas ÷ VCPs programadas"),
    ("13", "Treinamentos Realizados/Prog.",     "≥ 95%",      "Nº realizados ÷ programados × 100"),
    ("14", "Inspeções de EPI",                  "100%/trim.", "Cobertura trimestral"),
    ("15", "Autuações de Órgãos Fiscaliz.",     "= 0",        "Nº absoluto"),
]

GLOSSARIO = sorted([
    ("ADC",        "Água Doce Captada",                                           "Indicador ambiental"),
    ("APR",        "Análise Preliminar de Riscos",                                "Documento de campo para atividades não rotineiras"),
    ("ASO",        "Atestado de Saúde Ocupacional",                               "Indicador de saúde — NR-7"),
    ("AVP",        "Ações Verificadas Planejadas",                                "Componente do IPAS"),
    ("AVR",        "Ações Verificadas Realizadas",                                "Componente do IPAS"),
    ("CAT",        "Comunicação de Acidente de Trabalho",                         "Emitida no 1º dia útil após o acidente — CLT / Lei 8.213"),
    ("CIPA",       "Comissão Interna de Prevenção de Acidentes",                  "NR-5"),
    ("DDS",        "Diálogo Diário de Segurança",                                 "Indicador proativo — frequência diária"),
    ("EFA",        "Empregados Fisicamente Ativos",                               "Indicador de saúde"),
    ("EPC",        "Equipamento de Proteção Coletiva",                            "Controle operacional — NR-6"),
    ("EPI",        "Equipamento de Proteção Individual",                          "NR-6"),
    ("FISPQ",      "Ficha de Informações de Segurança de Produtos Químicos",      "NBR 14725 / ABNT"),
    ("HHER",       "Homens-Horas de Exposição ao Risco",                          "Base de cálculo de todos os indicadores de frequência"),
    ("HHERPR",     "HHER em Processos de Risco",                                  "Base de cálculo dos indicadores TASP_T1 e TASP_T2"),
    ("HP",         "Horas Programadas",                                           "Componente do PTP_S"),
    ("HS",         "Horas de Saúde Perdidas",                                     "Componente do PTP_S"),
    ("IAB",        "Índice de Implementação de Ações de Abrangência",             "Indicador proativo — trimestral"),
    ("IARAR",      "Índice de Acompanhamento das Recomendações de Análise de Riscos", "Indicador proativo — mensal"),
    ("IPAQ",       "International Physical Activity Questionnaire",               "Questionário internacional de atividade física — base do EFA"),
    ("IPAS",       "Índice de Plano de Ação Pós-Simulado",                        "Indicador de contingência — anual"),
    ("IPER",       "Identificação de Perigos e Avaliação de Riscos",              "Ferramenta PLAN — ISO 45001 §6"),
    ("IRS",        "Índice de Realização de Simulados",                           "Indicador de contingência — anual"),
    ("IRREEA",     "Índice de Resposta das Equipes de Evacuação Aeromédica",      "Indicador de saúde — unidades marítimas"),
    ("KML",        "Quilômetros Percorridos por Veículos Leves",                  "Componente do TFAT"),
    ("KMP",        "Quilômetros Percorridos por Veículos Pesados",                "Componente do TFAT"),
    ("LAIA",       "Levantamento de Aspectos e Impactos Ambientais",              "Ferramenta PLAN — ISO 14001 §6"),
    ("LOTO",       "Lockout / Tagout — Isolamento e Bloqueio de Energias",        "NR-10 — controle operacional"),
    ("NAF",        "Número de Acidentados Fatais",                                "Componente do TAF"),
    ("NAL",        "Número de Acidentados em Veículos Leves",                     "Componente do TFAT"),
    ("NAP",        "Número de Anomalias com Alto Potencial",                      "Componente do TAP"),
    ("NAPt",       "Número de Acidentados em Veículos Pesados",                   "Componente do TFAT — diferente do NAP de alto potencial"),
    ("NAR",        "Número de Acidentados Registráveis",                          "Exclui Classe 1 (primeiros socorros) no TAR"),
    ("NAR₄₅",     "Número de Acidentados Registráveis — Classes 4 e 5",          "Componente do TAG"),
    ("NARC",       "Número de Análises de Riscos em Conformidade",                "Componente do PARC"),
    ("NASP_T1",    "Número de Anomalias de Segurança de Processo — Nível 1",      "Componente do TASP_T1"),
    ("NASP_T2",    "Número de Anomalias de Segurança de Processo — Nível 2",      "Componente do TASP_T2"),
    ("NASOV",      "Número de ASOs Válidos",                                      "Componente do indicador ASO"),
    ("NCA",        "Número de Acidentados Com Afastamento",                       "Componente do TFCA"),
    ("NEA",        "Número de Empregados Ativos",                                 "Base do indicador ASO"),
    ("NR",         "Norma Regulamentadora",                                       "Legislação trabalhista — MTE / SESMT"),
    ("NRAPO",      "Nº de Recomendações Aprovadas em Prazo Original",             "Componente do IARAR"),
    ("NRAPP",      "Nº de Recomendações Aprovadas com Prazo Prorrogado",          "Componente do IARAR"),
    ("NREI",       "Nº de Recomendações Encerradas Implantadas",                  "Componente do IARAR"),
    ("NRNJ",       "Nº de Recomendações Não Justificadas",                        "Componente do IARAR"),
    ("NSA",        "Número de Acidentados Sem Afastamento",                       "Componente do TFSA"),
    ("NTAR",       "Número Total de Análises de Riscos",                          "Componente do PARC"),
    ("NTR",        "Número Total de Recomendações",                               "Componente do IARAR"),
    ("PAE",        "Plano de Ação de Emergência",                                 "Ciclo DO — resposta a emergências"),
    ("PARC",       "Percentual de Análises de Riscos em Conformidade",            "Indicador proativo — trimestral"),
    ("PEX",        "Procedimento de Execução",                                    "Documento técnico para atividades críticas"),
    ("PGRS",       "Plano de Gerenciamento de Resíduos Sólidos",                  "Gestão ambiental — ciclo DO"),
    ("PRE",        "Plano de Resposta a Emergências",                             "Ciclo DO — §2.6"),
    ("PS",         "Primeiros Socorros — Classe 1",                               "Não computado no TAR"),
    ("PT",         "Permissão de Trabalho",                                       "Controle operacional — atividades de alto risco"),
    ("PTP_S",      "Percentual de Tempo Perdido Saúde",                           "Indicador de saúde"),
    ("REM",        "Relatório Estatístico Mensal",                                "Entrega à Petrobras até o 5º dia do mês seguinte"),
    ("RRR",        "Reciclagem, Recuperação e Reúso",                             "Rotas preferenciais para destinação de resíduos sólidos"),
    ("RSNPA",      "Resíduos Sólidos Não Perigosos Acumulados",                   "Indicador ambiental"),
    ("RSNPD_NRRR", "Resíduos Sólidos NP Destinados — rotas Não-RRR",              "Indicador ambiental"),
    ("RSNPG",      "Resíduos Sólidos Não Perigosos Gerados",                      "Indicador ambiental"),
    ("RSPA",       "Resíduos Sólidos Perigosos Acumulados",                       "Indicador ambiental"),
    ("RSPD_NRRR",  "Resíduos Sólidos Perigosos Destinados — rotas Não-RRR",       "Indicador ambiental"),
    ("RSPG",       "Resíduos Sólidos Perigosos Gerados",                          "Indicador ambiental"),
    ("SMS",        "Segurança, Meio Ambiente e Saúde",                            "Sistema de gestão integrado — ISO 14001 + ISO 45001"),
    ("SP_C",       "Simulados Planejados por Contrato",                           "Componente do IRS"),
    ("SR_C",       "Simulados Realizados por Contrato",                           "Componente do IRS"),
    ("TAG",        "Taxa de Acidentados Graves",                                  "Indicador reativo — Classes 4 e 5"),
    ("TAF",        "Taxa de Acidentados Fatais",                                  "Indicador reativo — fator 10⁸"),
    ("TAP",        "Taxa de Anomalias com Alto Potencial",                        "Indicador reativo"),
    ("TAR",        "Taxa de Acidentados Registráveis",                            "Indicador reativo — exclui Classe 1"),
    ("TASP",       "Taxa de Anomalias de Segurança de Processo",                  "Indicadores proativos — Nível 1 (T1) e Nível 2 (T2)"),
    ("TCA",        "Total de Dias Computados como Afastamento",                   "Componente da Taxa de Gravidade (TG)"),
    ("TFAT",       "Taxa de Frequência de Acidentes de Trânsito",                 "Indicador reativo"),
    ("TFCA",       "Taxa de Frequência de Acidentados Com Afastamento",           "Indicador reativo prioritário — REM"),
    ("TFSA",       "Taxa de Frequência de Acidentados Sem Afastamento",           "Indicador reativo — também chamado 'Tussas'"),
    ("TG",         "Taxa de Gravidade",                                           "Indicador reativo — trimestral"),
    ("TIDO",       "Taxa de Incidência de Doença Ocupacional",                    "Indicador de saúde"),
    ("TOR",        "Taxa de Ocorrências Registráveis",                            "Indicador reativo"),
    ("VAR",        "Volume de Água Reusada",                                      "Indicador ambiental"),
    ("VAZO",       "Volume Vazado de Óleo e Derivados",                           "Volumes > 1 barril que atingiram corpos hídricos ou solo"),
    ("VCP",        "Verificação Comportamental Preventiva",                       "Indicador proativo — frequência mensal"),
], key=lambda x: x[0])

# ── Documento ──────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

# Cabeçalho
hdr = doc.add_table(rows=1, cols=2)
hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr.columns[0].width = Cm(11.5)
hdr.columns[1].width = Cm(6)
set_bg(hdr.cell(0, 0), "0F2D52")
set_bg(hdr.cell(0, 1), "E07B2A")
hdr.cell(0, 0).text = ""
p0 = hdr.cell(0, 0).paragraphs[0]
p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p0.add_run("JFX ENGENHARIA  |  GESTÃO SMS\n")
r.bold = True; r.font.name = FONT; r.font.size = Pt(13); r.font.color.rgb = WHITE
r = p0.add_run("EDISER / Petrobras — Blocos B, D e Q — ICJ 5900.0132378.25.2")
r.font.name = FONT; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0xCC, 0xD9, 0xEA)
hdr.cell(0, 1).text = ""
p1 = hdr.cell(0, 1).paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p1.add_run("INDICADORES SMS\n")
r.bold = True; r.font.name = FONT; r.font.size = Pt(12); r.font.color.rgb = WHITE
r = p1.add_run("Fórmulas · Metas · Glossário")
r.font.name = FONT; r.font.size = Pt(9); r.font.color.rgb = WHITE

doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(2)
r = tp.add_run("INDICADORES DE GESTÃO SMS — JFX / EDISER / PETROBRAS")
r.bold = True; r.font.name = FONT; r.font.size = Pt(14); r.font.color.rgb = NAVY
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(6)
r = sub.add_run("Fonte: Indicadores Gestão SMS — Petrobras · PGS-JFX-EDISER-001 Rev 00 · ISO 45001 · ISO 14001 · NBR 14280")
r.italic = True; r.font.name = FONT; r.font.size = Pt(8.5); r.font.color.rgb = GRAY
divider(doc)

# ── 1. HHER ────────────────────────────────────────────────
add_section_title(doc, "1. BASE DE CÁLCULO — HHER")
divider(doc, "E07B2A")
make_table(doc,
    headers=["Sigla", "Nome Completo", "Definição"],
    rows_data=[[
        ("HHER",  {"bold": True}),
        ("Homens-Horas de Exposição ao Risco", {}),
        ("Somatório das horas trabalhadas por todos os empregados e contratados, incluindo horas extras. "
         "É a base obrigatória de todos os indicadores de frequência.", {}),
    ]],
    col_widths=[2, 4.5, 11],
)
add_note(doc,
    "Regime obras/campo (JFX): 144 h/mês por colaborador.  "
    "Para o REM mensal: usar média móvel dos últimos 14 meses, desconsiderando os 2 meses mais recentes.  "
    "No fechamento anual: usar valor real apurado.  "
    "Responsável JFX: Roberta Muniz (TST Adm).")

# ── 2. SEGURANÇA — REATIVOS ─────────────────────────────────
add_section_title(doc, "2. INDICADORES DE SEGURANÇA")
add_subsection(doc, "2.1  Indicadores Reativos  (frequência mensal, exceto TG: trimestral)")
divider(doc, "E07B2A")

rows = []
for sigla, nome, formula, meta, freq, melhor, obs in IND_SEG_REATIVO:
    nome_cell = nome + (f"  —  {obs}" if obs else "")
    rows.append([
        (sigla,     {"bold": True, "color": NAVY}),
        (nome_cell, {}),
        (formula,   {"italic": True}),
        (meta,      {"bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (freq,      {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (melhor,    {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
    ])
make_table(doc,
    headers=["Sigla", "Nome / Observação", "Fórmula", "Meta", "Freq.", "Melhor"],
    rows_data=rows, col_widths=[1.8, 5.8, 4.2, 1.5, 1.8, 1.5])

add_note(doc,
    "Classes de acidente (NBR 14280 + Petrobras):  "
    "1 = Primeiros socorros  |  2 = Lesão sem afastamento  |  "
    "3 = Lesão com afastamento temporário  |  "
    "4 = Incapacidade permanente parcial / Doença ocupacional  |  "
    "5 = Morte / Incapacidade permanente total.  "
    "Não computar: acidentes de trajeto e ocorrências não apropriáveis (Anexo E – PP-1PB-00150).")

# ── 2.2 PROATIVOS ───────────────────────────────────────────
add_subsection(doc, "2.2  Indicadores Proativos")
divider(doc, "E07B2A")

rows = []
for sigla, nome, formula, meta, freq, melhor, obs in IND_SEG_PROATIVO:
    nome_cell = nome + (f"  —  {obs}" if obs else "")
    rows.append([
        (sigla,     {"bold": True, "color": NAVY}),
        (nome_cell, {}),
        (formula,   {"italic": True}),
        (meta,      {"bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (freq,      {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (melhor,    {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
    ])
make_table(doc,
    headers=["Sigla", "Nome / Observação", "Fórmula", "Meta", "Freq.", "Melhor"],
    rows_data=rows, col_widths=[1.8, 5.8, 4.2, 1.5, 1.8, 1.5])

# ── 3. CONTINGÊNCIA ─────────────────────────────────────────
add_section_title(doc, "3. INDICADORES DE CONTINGÊNCIA")
divider(doc, "E07B2A")

rows = []
for sigla, nome, formula, meta, freq, melhor, obs in IND_CONTINGENCIA:
    rows.append([
        (sigla,   {"bold": True, "color": NAVY}),
        (nome,    {}),
        (formula, {"italic": True}),
        (meta,    {"bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (freq,    {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (melhor,  {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (obs,     {"size": 8.5}),
    ])
make_table(doc,
    headers=["Sigla", "Nome", "Fórmula", "Meta", "Freq.", "Melhor", "Observação"],
    rows_data=rows, col_widths=[1.5, 4.0, 3.2, 1.2, 1.5, 1.5, 4.7])

# ── 4. AMBIENTAL ────────────────────────────────────────────
add_section_title(doc, "4. INDICADORES AMBIENTAIS")
divider(doc, "E07B2A")

rows = []
for sigla, nome, apuracao, unid, freq, melhor, obs in IND_AMBIENTAL:
    nome_cell = nome + (f"  —  {obs}" if obs else "")
    rows.append([
        (sigla,     {"bold": True, "color": NAVY}),
        (nome_cell, {}),
        (unid,      {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (apuracao,  {"italic": True}),
        (freq,      {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (melhor,    {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
    ])
make_table(doc,
    headers=["Sigla", "Nome / Observação", "Unid.", "Apuração", "Freq.", "Melhor"],
    rows_data=rows, col_widths=[2.2, 7.3, 1.2, 3.5, 1.5, 1.5])
add_note(doc, "RRR = Reciclagem, Recuperação e Reúso — rotas preferenciais de destinação de resíduos sólidos.")

# ── 5. SAÚDE ────────────────────────────────────────────────
add_section_title(doc, "5. INDICADORES DE SAÚDE")
divider(doc, "E07B2A")

rows = []
for sigla, nome, formula, melhor, freq, obs in IND_SAUDE:
    rows.append([
        (sigla,   {"bold": True, "color": NAVY}),
        (nome,    {}),
        (formula, {"italic": True}),
        (melhor,  {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (freq,    {"size": 8.5, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (obs,     {"size": 8.5}),
    ])
make_table(doc,
    headers=["Sigla", "Nome", "Fórmula", "Melhor", "Freq.", "Observação"],
    rows_data=rows, col_widths=[1.8, 4.5, 4.2, 1.5, 1.5, 4.2])

# ── 6. REM JFX ─────────────────────────────────────────────
add_section_title(doc, "6. INDICADORES PRIORITÁRIOS — REM JFX / EDISER  (15 indicadores)")
divider(doc, "E07B2A")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run(
    "Indicadores exigidos no REM mensal do contrato EDISER (PGS-JFX-EDISER-001, item 5.1).  "
    "Prazo: até o 5º dia do mês subsequente.  "
    "Responsável: Roberta Muniz (TST Adm).  Validação: Fábio Martins Nascimento (Coord. SMS)."
)
r.font.name = FONT; r.font.size = Pt(9)

rows = []
for num, ind, meta, formula in REM_15:
    rows.append([
        (num,     {"bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (ind,     {"bold": True, "color": NAVY}),
        (meta,    {"bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER}),
        (formula, {"italic": True}),
    ])
make_table(doc,
    headers=["#", "Indicador", "Meta", "Base de Cálculo / Fórmula"],
    rows_data=rows, col_widths=[0.8, 5.5, 2.0, 9.3])

# ── 7. GLOSSÁRIO ────────────────────────────────────────────
add_section_title(doc, "7. GLOSSÁRIO DE SIGLAS  (ordem alfabética)")
divider(doc, "E07B2A")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Todas as siglas utilizadas neste documento com respectiva expansão e contexto de uso.")
r.font.name = FONT; r.font.size = Pt(9)

rows = []
for sigla, expansao, contexto in GLOSSARIO:
    rows.append([
        (sigla,    {"bold": True, "color": NAVY}),
        (expansao, {}),
        (contexto, {"italic": True, "color": GRAY}),
    ])
make_table(doc,
    headers=["Sigla", "Expansão", "Contexto / Observação"],
    rows_data=rows, col_widths=[2.8, 8.5, 6.3])

aplicar_rodape(doc,
    linha_extra="Documento gerado em 03/05/2026  ·  Uso interno  ·  "
                "Ref.: ISO 45001 · ISO 14001 · NBR 14280 · N-Petrobras"
)

# ── Salvar ──────────────────────────────────────────────────
out = r"C:\Users\fabio\Documents\segundo-cerebro\Indicadores-SMS-JFX-EDISER.docx"
doc.save(out)
print(f"OK: {out}")
