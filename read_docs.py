import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
import os

# Read the old PEX
p1 = r"C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\pex\PEX - BLOCO B - EXECUÇÃO DE BANCADAS EM GRANITO.docx"
print("=== OLD PEX ===")
doc = Document(p1)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[P{i}] {para.style.name!r} | {para.text[:200]}")
print("--- TABLES ---")
for ti, table in enumerate(doc.tables):
    print(f"Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip()[:80] for c in row.cells]
        print(f"  R{ri}: {cells}")

# Read reference PEX (Forro Rev_C)
p2 = r"C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\APRs - EDISER\PEX-JFX-ENG-033-INSTALACAO-FORRO-Rev_C.docx"
print("\n\n=== REFERENCE PEX (FORRO Rev_C) ===")
doc2 = Document(p2)
for i, para in enumerate(doc2.paragraphs):
    if para.text.strip():
        print(f"[P{i}] {para.style.name!r} | {para.text[:200]}")
print("--- TABLES ---")
for ti, table in enumerate(doc2.tables):
    print(f"Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip()[:80] for c in row.cells]
        print(f"  R{ri}: {cells}")
