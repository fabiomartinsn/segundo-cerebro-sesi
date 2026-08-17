from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

NAVY  = RGBColor(0x0F, 0x2D, 0x52)
LARAN = RGBColor(0xE0, 0x7B, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CINZA = RGBColor(0xF2, 0xF2, 0xF2)
VERDE = RGBColor(0x1A, 0x7A, 0x1A)
VERM  = RGBColor(0xC0, 0x39, 0x2B)
AMAR  = RGBColor(0xF3, 0x9C, 0x12)

def set_cell_bg(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = LARAN
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    # Borda inferior nos headings principais
    if level in (1, 2):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),  'single')
        bottom.set(qn('w:sz'),   '6')
        bottom.set(qn('w:space'),'1')
        bottom.set(qn('w:color'), '0F2D52' if level == 1 else 'E07B2A')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(doc, text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    return p

def table_header_row(table, headers, bg='0F2D52'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        run.font.size = Pt(9)

def add_table_row(table, values, bg=None, bold=False, colors=None):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg[i] if isinstance(bg, list) else bg)
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.bold = bold
        if colors and colors[i]:
            run.font.color.rgb = colors[i]
    return row

# ── DOCUMENTO ──────────────────────────────────────────────────────────────

doc = Document()

# Margens
sec = doc.sections[0]
sec.top_margin    = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

# Estilo padrão
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── CABEÇALHO ──────────────────────────────────────────────────────────────
tbl = doc.add_table(rows=2, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.columns[0].width = Cm(12)
tbl.columns[1].width = Cm(6.5)

set_cell_bg(tbl.cell(0,0), '0F2D52')
c00 = tbl.cell(0,0).paragraphs[0]
r = c00.add_run('JFX ENGENHARIA')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = WHITE; r.font.name = 'Calibri'

set_cell_bg(tbl.cell(0,1), '0F2D52')
c01 = tbl.cell(0,1).paragraphs[0]
c01.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = c01.add_run(f'Emitido em: {datetime.date.today().strftime("%d/%m/%Y")}')
r.font.size = Pt(9); r.font.color.rgb = WHITE; r.font.name = 'Calibri'

set_cell_bg(tbl.cell(1,0), 'E07B2A')
c10 = tbl.cell(1,0).paragraphs[0]
r = c10.add_run('PLANO DE DOCUMENTAÇÃO — ÁREA LIBERADA BLOCO D  |  EDISER / Petrobras')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE; r.font.name = 'Calibri'

set_cell_bg(tbl.cell(1,1), 'E07B2A')
c11 = tbl.cell(1,1).paragraphs[0]
c11.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = c11.add_run('ICJ 5900.0132378.25.2\nCoordenador SMS: Fábio Martins Nascimento\nCREA-SE 2718664207')
r.font.size = Pt(8); r.font.color.rgb = WHITE; r.font.name = 'Calibri'

doc.add_paragraph()

# ── 1. REFERÊNCIA ──────────────────────────────────────────────────────────
add_heading(doc, '1. REFERÊNCIA E OBJETIVO')
add_body(doc, (
    'Este documento foi elaborado com base na Área Liberada N°002/2026 — Bloco B (PE-2ACO-00005 / Anexo E), '
    'vigência 16/03/2026 a 30/12/2026, adotada como modelo padrão para a Área Liberada do Bloco D. '
    'O objetivo é mapear todos os documentos necessários, identificar o que já está disponível, '
    'o que precisa ser produzido e definir os próximos passos para protocolização.'
))

# ── 2. DOCUMENTOS ESTRUTURAIS ──────────────────────────────────────────────
add_heading(doc, '2. DOCUMENTOS ESTRUTURAIS')
add_body(doc, 'Documentos que compõem o corpo principal da Área Liberada, independente dos colaboradores.', color=NAVY)

doc.add_paragraph()
t1 = doc.add_table(rows=1, cols=4)
t1.style = 'Table Grid'
table_header_row(t1, ['#', 'Documento', 'Status', 'Observação'])
t1.columns[0].width = Cm(0.8)
t1.columns[1].width = Cm(7)
t1.columns[2].width = Cm(2.5)
t1.columns[3].width = Cm(8.2)

estruturais = [
    ('1', 'Formulário Anexo E (PE-2ACO-00005) — Bloco D',
     '❌ FALTA', 'Novo formulário a ser preenchido para o Bloco D. Seguir modelo do Bloco B.'),
    ('2', 'Planta baixa / croqui da área de intervenção — Bloco D',
     '❌ FALTA', 'Indicar salas, pavimentos e perímetro exato de intervenção conforme projeto.'),
    ('3', 'APRs de todas as atividades previstas no Bloco D (mín. 22)',
     '⚠️ ADAPTAR', 'As 22 APRs do Bloco B existem. Revisar escopo e adaptar para o Bloco D. Adicionar atividades novas se houver.'),
    ('4', 'PEXes de todas as atividades (mín. 19)',
     '⚠️ ADAPTAR', 'Os 19 PEXes do Bloco B existem. Adaptar numeração e escopo para o Bloco D. Adendos N-2869 já aplicados.'),
    ('5', 'Folha diária de liberação (Premissa 9)',
     '✅ MODELO OK', 'Modelo do Bloco B em uso. Gerar folha nova com cabeçalho Bloco D.'),
]

for row in estruturais:
    st = row[2]
    cor = 'F2F2F2' if '✅' in st else ('FFF3CD' if '⚠️' in st else 'FDECEA')
    add_table_row(t1, row, bg=cor)

doc.add_paragraph()

# ── 3. PROGRAMAS DE SMS ────────────────────────────────────────────────────
add_heading(doc, '3. PROGRAMAS DE SMS (Premissa 1)')
add_body(doc, 'Todos os 9 programas exigidos já existem e estão arquivados em vault/_knowledge/Petrobras/Area Liberada/PROGRAMAS DE SMS/. '
              'Verificar validade/revisão antes de protocolar para o Bloco D.')

doc.add_paragraph()
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
table_header_row(t2, ['#', 'Documento', 'Código', 'Status'])

programas = [
    ('1', 'Programa de Gerenciamento de Riscos',             'PGR',        '✅ Disponível'),
    ('2', 'Programa de Controle Médico de Saúde Ocupacional','PCMSO',      '✅ Disponível'),
    ('3', 'Programa de Conservação Auditiva',                'PGR-SMS002 / PCA',  '✅ Disponível'),
    ('4', 'Programa de Proteção Respiratória',               'PGR-SMS003 / PPR',  '✅ Disponível'),
    ('5', 'Análise Ergonômica do Trabalho',                  'PGR-SMS005 / AET',  '✅ Disponível'),
    ('6', 'Plano de Atendimento a Emergências',              'PL-SMS-002 / PAE',  '✅ Disponível'),
    ('7', 'Plano de Emergência Médica / Primeiros Socorros', 'PL-SMS-003 / PEMPS','✅ Disponível'),
    ('8', 'Plano de Resposta a Emergências',                 'PL-SMS-004 / PRE',  '✅ Disponível'),
    ('9', 'Plano de Gerenciamento de Resíduos da Construção Civil', 'PL-SGI-002 / PGRS', '✅ Disponível'),
]

for row in programas:
    add_table_row(t2, row, bg='E8F5E9')

doc.add_paragraph()

# ── 4. DOCUMENTAÇÃO DOS COLABORADORES ─────────────────────────────────────
add_heading(doc, '4. DOCUMENTAÇÃO POR COLABORADOR')
add_body(doc, (
    'Para cada colaborador alocado no Bloco D, os documentos abaixo são obrigatórios conforme '
    'Premissa 12 da Área Liberada. A lista de colaboradores do Bloco D ainda não está definida — '
    'preencher antes de protocolar o processo.'
))

doc.add_paragraph()
t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
table_header_row(t3, ['#', 'Documento Exigido', 'Observação'])

docs_colab = [
    ('1',  'Certificado NR-06 — EPIs',              'Todos os colaboradores'),
    ('2',  'Certificado NR-10 — Elétrica',          'Obrigatório para funções que envolvam elétrica e TST'),
    ('3',  'Certificado NR-11 — Movimentação de Cargas', 'Todos os colaboradores'),
    ('4',  'Certificado NR-12 — Máquinas e Equipamentos', 'Todos os colaboradores'),
    ('5',  'Certificado NR-17 — Ergonomia',         'Todos os colaboradores'),
    ('6',  'Certificado NR-18 — Construção Civil',  'Todos os colaboradores'),
    ('7',  'Certificado NR-35 — Trabalho em Altura','Todos os colaboradores'),
    ('8',  'Ordem de Serviço (OS) assinada',        'Específica para cada função — template disponível'),
    ('9',  'Ficha de EPI assinada',                 'Registro de entrega de todos os EPIs individuais'),
    ('10', 'Comprovante eSocial (S-2220 / S-2240)', 'Obrigatório conforme exigência Petrobras'),
    ('11', 'ASO — Atestado de Saúde Ocupacional',   'Válido — verificar vencimento antes de protocolar'),
    ('12', 'Regras de Ouro assinadas',               'Declaração individual de ciência das Regras de Ouro'),
    ('13', 'Registro de treinamento de Área Liberada (FOR-SGI 001)', 'Mínimo 30 min de treinamento + lista de presença assinada'),
]

for i, row in enumerate(docs_colab):
    bg = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
    add_table_row(t3, row, bg=bg)

doc.add_paragraph()

# ── 5. COLABORADORES BLOCO D ───────────────────────────────────────────────
add_heading(doc, '5. RELAÇÃO DE COLABORADORES — BLOCO D')
add_body(doc, '⚠️ Lista a ser preenchida. Inserir todos os colaboradores que atuarão no Bloco D antes da protocolização.', color=VERM, bold=True)

doc.add_paragraph()
t4 = doc.add_table(rows=1, cols=14)
t4.style = 'Table Grid'
table_header_row(t4, ['#','Colaborador','Função','NR06','NR10','NR11','NR12','NR17','NR18','NR35','OS','EPI','eSocial','ASO'])

for i in range(1, 9):
    row = t4.add_row()
    row.cells[0].paragraphs[0].add_run(str(i)).font.size = Pt(8)
    for j in range(1, 14):
        row.cells[j].paragraphs[0].add_run('').font.size = Pt(8)
    set_cell_bg(row.cells[0], 'F2F2F2')

doc.add_paragraph()
add_body(doc, 'Legenda:  ✅ Presente  |  ❌ Ausente  |  ⚠️ Vencido / Pendente de revisão', color=NAVY)
doc.add_paragraph()

# ── 6. ATIVIDADES E APRs/PEXes ────────────────────────────────────────────
add_heading(doc, '6. ATIVIDADES PREVISTAS — APRs E PEXes NECESSÁRIOS')
add_body(doc, (
    'Confirmar escopo de atividades do Bloco D com a equipe de engenharia. As APRs e PEXes do Bloco B '
    'servem como base; adaptar para o Bloco D e submeter à aprovação prévia da fiscalização Petrobras (Premissa 4).'
))

doc.add_paragraph()
t5 = doc.add_table(rows=1, cols=5)
t5.style = 'Table Grid'
table_header_row(t5, ['APR Bloco B (Referência)', 'PEX Bloco B (Referência)', 'Prevista Bloco D?', 'APR Adaptada', 'PEX Adaptado'])

atividades_ref = [
    ('APR-001 Retirada de Mobiliário',             'PEX-JFX-ENG-001 Demolições e Retiradas',          '','',''),
    ('APR-002 Retirada de Vidros e Divisórias',     '—',                                                '','',''),
    ('APR-003 Demolição de Forros',                 '—',                                                '','',''),
    ('APR-004 Demolição de Alvenaria e Rodapés',    '—',                                                '','',''),
    ('APR-005 Alvenaria de Vedação',                'PEX-JFX-ENG-002 Alvenaria de Vedação',            '','',''),
    ('APR-006 Aplicação de Carpete Boucle',         'PEX-JFX-ENG-003 Carpete Boucle',                  '','',''),
    ('APR-007 Execução de Concreto',                'PEX-JFX-ENG-004 Concreto Usinado',                '','',''),
    ('APR-008 Montagem e Fixação de Ferragem',      'PEX-JFX-ENG-005 Ferragem',                        '','',''),
    ('APR-009 Corte e Montagem de Forma de Concreto','PEX-JFX-ENG-006 Forma de Concreto',              '','',''),
    ('APR-010 Revestimento Cerâmico',               'PEX-JFX-ENG-007 Revestimento Cerâmico',           '','',''),
    ('APR-011 Revestimento de Argamassa',           'PEX-JFX-ENG-008 Argamassa',                       '','',''),
    ('APR-012 Revestimento Vinílico e Rodapé',      'PEX-JFX-ENG-009 Vinílico e Rodapé',               '','',''),
    ('APR-013 Instalação de Esquadrias',            'PEX-JFX-ENG-010 Esquadrias',                      '','',''),
    ('APR-014 Louças e Metais',                     'PEX-JFX-ENG-011 Louças e Metais Sanitários',      '','',''),
    ('APR-015 Instalações Hidrossanitárias',        'PEX-JFX-ENG-012 Hidrossanitárias',                '','',''),
    ('APR-016 Remoção de Ar-Condicionado',          'PEX-JFX-ENG-013 Remoção AC e Duto',               '','',''),
    ('APR-017 HVAC — Dutos e Difusores',            'PEX-JFX-ENG-014 HVAC Dutos',                      '','',''),
    ('APR-018 HVAC — VRF e Ventilação Mecânica',    'PEX-JFX-ENG-015 HVAC VRF',                        '','',''),
    ('APR-019 Cabeamento Estruturado',              'PEX-JFX-ENG-016 Cabeamento',                      '','',''),
    ('APR-020 SPCI — Incêndio',                    'PEX-JFX-ENG-017 SPCI',                            '','',''),
    ('APR-021 Elétrica Iluminação/Tomadas',         'PEX-JFX-ENG-018 Elétrica (infra)',                 '','',''),
    ('APR-022 Quadros Elétricos',                   'PEX-JFX-ENG-019 Quadros Elétricos',               '','',''),
]

for i, row in enumerate(atividades_ref):
    bg = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
    add_table_row(t5, row, bg=bg)

doc.add_paragraph()

# ── 7. O QUE FALTA — RESUMO EXECUTIVO ─────────────────────────────────────
add_heading(doc, '7. RESUMO EXECUTIVO — O QUE FALTA')

doc.add_paragraph()
t6 = doc.add_table(rows=1, cols=4)
t6.style = 'Table Grid'
table_header_row(t6, ['Categoria', 'Item', 'Situação', 'Responsável'])

pendencias = [
    ('Formulário Oficial',    'Anexo E — Bloco D (PE-2ACO-00005)',                         '❌ FALTA',   'Fábio Martins + Lúcio'),
    ('Planta / Croqui',       'Planta baixa do Bloco D com área de intervenção demarcada', '❌ FALTA',   'Eng. Responsável'),
    ('Colaboradores',         'Definir lista completa de colaboradores do Bloco D',         '❌ FALTA',   'Preposto Lúcio'),
    ('Docs por Colaborador',  '13 documentos por colaborador (NRs, OS, EPI, ASO, eSocial, Reg. Ouro, Trein. AL)', '❌ FALTA', 'Roberta Muniz'),
    ('APRs',                  '22 APRs a adaptar do Bloco B para o Bloco D',               '⚠️ ADAPTAR', 'Fábio Martins'),
    ('PEXes',                 '19 PEXes a adaptar + adendos N-2869 aplicar ao Bloco D',    '⚠️ ADAPTAR', 'Fábio Martins + Eng. Lúcio'),
    ('Treinamento AL',        'Realizar treinamento de Área Liberada com equipe Bloco D',  '❌ FALTA',   'Fábio + Roberta'),
    ('Folha Diária',          'Folha diária de liberação com cabeçalho Bloco D',           '⚠️ ADAPTAR', 'Fábio Martins'),
    ('Programas SMS',         '9 programas SMS (PGR, PCMSO, PCA, PPR, AET, PAE, PEMPS, PRE, PGRS)', '✅ DISPONÍVEL', '—'),
]

for row in pendencias:
    st = row[2]
    if '✅' in st:  bg = 'E8F5E9'
    elif '⚠️' in st: bg = 'FFF3CD'
    else:            bg = 'FDECEA'
    add_table_row(t6, row, bg=bg)

doc.add_paragraph()

# ── 8. PRÓXIMOS PASSOS ─────────────────────────────────────────────────────
add_heading(doc, '8. PRÓXIMOS PASSOS')

passos = [
    ('1', 'IMEDIATO', 'Definir lista de colaboradores do Bloco D com o Preposto Lúcio.'),
    ('2', 'IMEDIATO', 'Confirmar escopo de atividades do Bloco D com engenharia — quais APRs/PEXes se aplicam.'),
    ('3', 'CURTO PRAZO', 'Adaptar APRs do Bloco B para o Bloco D — revisar riscos específicos e responsáveis.'),
    ('4', 'CURTO PRAZO', 'Adaptar PEXes do Bloco B + aplicar adendos N-2869 ao Bloco D.'),
    ('5', 'CURTO PRAZO', 'Coletar documentação individual de cada colaborador (13 itens por pessoa).'),
    ('6', 'CURTO PRAZO', 'Realizar treinamento de Área Liberada com equipe Bloco D (mín. 30 min) — gerar lista de presença assinada.'),
    ('7', 'ANTES DO PROTOCOLO', 'Preencher Formulário Anexo E (PE-2ACO-00005) com dados do Bloco D.'),
    ('8', 'ANTES DO PROTOCOLO', 'Obter croqui / planta da área de intervenção Bloco D.'),
    ('9', 'PROTOCOLO', 'Encaminhar dossiê completo à fiscalização Petrobras para análise e assinatura (Regina Celia / Daniel Machado / Vitor Pinheiro / Diego Pimentel / Priscilla Menegatti).'),
    ('10','APÓS APROVAÇÃO', 'Manter toda a documentação na frente de obras conforme Premissa 1. Realizar folha diária de liberação (Premissa 9).'),
]

doc.add_paragraph()
t7 = doc.add_table(rows=1, cols=4)
t7.style = 'Table Grid'
table_header_row(t7, ['#', 'Prazo', 'Ação', 'Responsável'])

for p in passos:
    if p[1] == 'IMEDIATO':        bg = 'FDECEA'
    elif p[1] == 'CURTO PRAZO':   bg = 'FFF3CD'
    elif 'PROTOCOLO' in p[1]:     bg = 'E3F2FD'
    else:                          bg = 'E8F5E9'
    add_table_row(t7, (p[0], p[1], p[2], '—'), bg=bg)

doc.add_paragraph()

# ── RODAPÉ ─────────────────────────────────────────────────────────────────
add_body(doc,
    'JFX Engenharia  |  CREA-SE 2718664207  |  Fábio Martins Nascimento — Coordenador de SMS  |  '
    f'Emitido em: {datetime.date.today().strftime("%d/%m/%Y")}  |  Gerado por FabIA — Segundo Cérebro',
    color=NAVY, size=8
)

# ── SALVAR ─────────────────────────────────────────────────────────────────
out = r'C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\AREA-LIBERADA-BLOCO-D-Plano-Documentacao.docx'
doc.save(out)
print(f'Salvo: {out}')
