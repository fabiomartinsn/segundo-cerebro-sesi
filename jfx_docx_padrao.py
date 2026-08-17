"""
Padrão JFX — rodapé e identidade visual para documentos .docx
Importar em qualquer script de geração:

    from jfx_docx_padrao import aplicar_rodape, NAVY, ORANGE, WHITE, GRAY, FONT
"""

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta JFX ─────────────────────────────────────────────
NAVY   = RGBColor(0x0F, 0x2D, 0x52)
ORANGE = RGBColor(0xE0, 0x7B, 0x2A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x77, 0x77, 0x77)
FONT   = "Calibri"

# ── Dados do responsável ───────────────────────────────────
RESPONSAVEL = "Eng. Fábio Martins"
CREA        = "CREA 2718664207"
CARGO       = "Coord. SMS"
EMPRESA     = "JFX Engenharia"
CONTRATO    = "ICJ 5900.0132378.25.2  |  EDISER / Petrobras"

def aplicar_rodape(doc, linha_extra: str = ""):
    """
    Insere o rodapé padrão JFX em todas as seções do documento.

    Linha 1: empresa + contrato
    Linha 2: responsável + CREA + cargo
    Linha extra (opcional): data, versão, observação etc.

    Uso:
        aplicar_rodape(doc)
        aplicar_rodape(doc, linha_extra="Documento gerado em 04/05/2026  ·  Uso interno")
    """
    for section in doc.sections:
        section.footer_distance = Cm(1)
        footer = section.footer
        # limpar parágrafos existentes
        for p in footer.paragraphs:
            p.clear()

        # ── separador ──────────────────────────────────────
        p_sep = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(2)
        p_sep.paragraph_format.space_after  = Pt(0)
        pPr = p_sep._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"),   "single")
        top.set(qn("w:sz"),    "4")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "0F2D52")
        pBdr.append(top)
        pPr.append(pBdr)

        # ── linha 1: empresa + contrato ────────────────────
        p1 = footer.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after  = Pt(0)
        r1 = p1.add_run(f"{EMPRESA}  |  {CONTRATO}")
        r1.font.name = FONT; r1.font.size = Pt(7.5)
        r1.font.color.rgb = GRAY

        # ── linha 2: responsável + CREA + cargo ───────────
        p2 = footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r2a = p2.add_run(f"{RESPONSAVEL}  —  {CREA}  —  {CARGO}")
        r2a.bold = True
        r2a.font.name = FONT; r2a.font.size = Pt(8)
        r2a.font.color.rgb = NAVY

        # ── linha extra (opcional) ─────────────────────────
        if linha_extra:
            p3 = footer.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after  = Pt(0)
            r3 = p3.add_run(linha_extra)
            r3.italic = True
            r3.font.name = FONT; r3.font.size = Pt(7.5)
            r3.font.color.rgb = GRAY
