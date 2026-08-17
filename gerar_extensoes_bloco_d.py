"""
Gera Folhas de Extensão de Aplicabilidade ao Bloco D
para todas as APRs e PEXes do contrato EDISER/JFX.
Padrão: mesmo formato das Folhas de Atualização N-2869.
"""
import os, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x0F, 0x2D, 0x52)
LARAN = RGBColor(0xE0, 0x7B, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUT_APR = r'C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\AREA-LIBERADA-BLOCO-D\APRs'
OUT_PEX = r'C:\Users\fabio\Documents\segundo-cerebro\vault\_knowledge\JFX\AREA-LIBERADA-BLOCO-D\PEXes'
DATA    = datetime.date.today().strftime('%d/%m/%Y')

# ── LISTA DE APRs ──────────────────────────────────────────────────────────
APRS = [
    ('APR-JFX-001', 'Rev.00', 'Retirada de Mobiliário Operacional'),
    ('APR-JFX-002', 'Rev.02', 'Retirada de Vidros e Divisórias'),
    ('APR-JFX-003', 'Rev.00', 'Demolição de Forros'),
    ('APR-JFX-004', 'Rev.00', 'Demolição de Alvenaria e Rodapés'),
    ('APR-JFX-005', 'Rev.00', 'Alvenaria de Vedação'),
    ('APR-JFX-006', 'Rev.00', 'Aplicação de Carpete Bouclê'),
    ('APR-JFX-007', 'Rev.00', 'Execução de Concreto'),
    ('APR-JFX-008', 'Rev.00', 'Montagem e Fixação de Ferragem'),
    ('APR-JFX-009', 'Rev.00', 'Corte e Montagem de Forma de Concreto'),
    ('APR-JFX-010', 'Rev.00', 'Revestimento Cerâmico de Piso e Parede'),
    ('APR-JFX-011', 'Rev.00', 'Revestimento de Argamassa'),
    ('APR-JFX-012', 'Rev.00', 'Revestimento Vinílico e Rodapé'),
    ('APR-JFX-013', 'Rev.00', 'Serviços de Instalações de Esquadrias'),
    ('APR-JFX-014', 'Rev.00', 'Remoção e Instalação de Louças e Metais Sanitários'),
    ('APR-JFX-015', 'Rev.00', 'Execução das Instalações Hidrossanitárias — Água Fria, Esgoto e Ventilação Predial'),
    ('APR-JFX-016', 'Rev.00', 'Serviços de Remoção de Ar-Condicionado e Duto'),
    ('APR-JFX-017', 'Rev.00', 'Sistema HVAC — Montagem de Dutos, Grelhas, Difusores, Venezianas e Dampers'),
    ('APR-JFX-018', 'Rev.00', 'Sistema HVAC (Sistemas VRF e Ventilação Mecânica)'),
    ('APR-JFX-019', 'Rev.00', 'Instalação de Cabeamento Estruturado para Redes de Dados e Voz'),
    ('APR-JFX-020', 'Rev.00', 'Instalações de SPCI — Sistema de Proteção Contra Incêndio'),
    ('APR-JFX-021', 'Rev.A',  'Instalação Elétrica de Iluminação e Tomadas (Desenergizado — Somente Infra)'),
    ('APR-JFX-022', 'Rev.00', 'Montagem e Instalação de Quadros Elétricos (Desenergizado — Somente Infra)'),
    ('APR-JFX-023', 'Rev.00', 'Uso de Betoneira Elétrica e Caminhão Betoneira'),
    ('APR-JFX-024', 'Rev.00', 'Instalação de Placa'),
    ('APR-JFX-ENG-020', 'Rev.01', 'Montagem, Utilização e Desmontagem de Andaimes — Internos'),
    ('APR-JFX-ENG-021', 'Rev.01', 'Montagem, Utilização e Desmontagem de Andaimes — Externos'),
    ('APR-JFX-ENG-025', 'Rev.00', 'Remoção de Piso Vinílico'),
    ('APR-JFX-ENG-026', 'Rev.01', 'Demolição de Piso'),
]

# ── LISTA DE PEXes ─────────────────────────────────────────────────────────
PEXES = [
    ('PEX-JFX-ENG-001', 'Rev.03', 'Demolições e Retiradas'),
    ('PEX-JFX-ENG-002', 'Rev.00', 'Alvenaria de Vedação'),
    ('PEX-JFX-ENG-003', 'Rev.00', 'Aplicação de Carpete Bouclê'),
    ('PEX-JFX-ENG-004', 'Rev.00', 'Execução de Concreto Usinado ou Dosado na Obra'),
    ('PEX-JFX-ENG-005', 'Rev.00', 'Montagem e Fixação de Ferragem'),
    ('PEX-JFX-ENG-006', 'Rev.00', 'Corte e Montagem de Forma de Concreto'),
    ('PEX-JFX-ENG-007', 'Rev.00', 'Revestimento Cerâmico de Piso e Parede'),
    ('PEX-JFX-ENG-008', 'Rev.00', 'Revestimento de Argamassa'),
    ('PEX-JFX-ENG-009', 'Rev.00', 'Revestimento Vinílico e Rodapé'),
    ('PEX-JFX-ENG-010', 'Rev.00', 'Serviços de Instalações de Esquadrias'),
    ('PEX-JFX-ENG-011', 'Rev.00', 'Remoção e Instalação de Louças e Metais Sanitários'),
    ('PEX-JFX-ENG-012', 'Rev.00', 'Execução das Instalações Hidrossanitárias — Água Fria, Esgoto e Ventilação Predial'),
    ('PEX-JFX-ENG-013', 'Rev.00', 'Serviços de Remoção de Ar-Condicionado e Duto'),
    ('PEX-JFX-ENG-014', 'Rev.00', 'Sistema HVAC — Montagem de Dutos, Grelhas, Difusores, Venezianas e Dampers'),
    ('PEX-JFX-ENG-015', 'Rev.00', 'Sistema HVAC (Sistemas VRF e Ventilação Mecânica)'),
    ('PEX-JFX-ENG-016', 'Rev.00', 'Instalação de Cabeamento Estruturado para Redes de Dados e Voz'),
    ('PEX-JFX-ENG-017', 'Rev.00', 'Instalações de SPCI — Sistema de Proteção Contra Incêndio'),
    ('PEX-JFX-ENG-018', 'Rev.A',  'Instalação Elétrica de Iluminação e Tomadas (Desenergizado — Somente Infra)'),
    ('PEX-JFX-ENG-019', 'Rev.00', 'Montagem e Instalação de Quadros Elétricos (Desenergizado — Somente Infra)'),
    ('PEX-JFX-ENG-020', 'Rev.01', 'Montagem, Utilização e Desmontagem de Andaimes — Internos'),
    ('PEX-JFX-ENG-021', 'Rev.01', 'Montagem, Utilização e Desmontagem de Andaimes — Externos'),
    ('PEX-JFX-ENG-023', 'Rev.00', 'Remoção de Piso Vinílico'),
    ('PEX-JFX-ENG-024', 'Rev.00', 'Demolição de Piso'),
]

# ── HELPERS ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_par(doc, text, bold=False, size=10, color=None, space_before=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def secao(doc, num, titulo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f'{num}  {titulo}')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = LARAN
    # borda inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'E07B2A')
    pBdr.append(bot); pPr.append(pBdr)

# ── GERADOR PRINCIPAL ──────────────────────────────────────────────────────
def gerar_folha(codigo, revisao, titulo, tipo, pasta_saida):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── CABEÇALHO ──
    t = doc.add_table(rows=3, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(10)
    t.columns[1].width = Cm(8)

    # linha 0
    set_cell_bg(t.cell(0,0), '0F2D52')
    cell_text(t.cell(0,0), 'JFX ENGENHARIA', bold=True, size=13, color=WHITE)
    set_cell_bg(t.cell(0,1), '0F2D52')
    cell_text(t.cell(0,1), 'FOLHA DE EXTENSÃO DE APLICABILIDADE — BLOCO D',
              bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # linha 1
    set_cell_bg(t.cell(1,0), 'E07B2A')
    cell_text(t.cell(1,0), f'{tipo} de Referência:  {codigo}', bold=True, size=10, color=WHITE)
    set_cell_bg(t.cell(1,1), 'E07B2A')
    cell_text(t.cell(1,1),
              f'N° do Documento:  EXT-{codigo}-BLOCO-D',
              bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # linha 2
    cell_text(t.cell(2,0), f'Título:  {titulo}', size=9)
    cell_text(t.cell(2,1),
              f'Rev. {tipo}: {revisao}  |  Data da Extensão: {DATA}',
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── SEÇÃO 1 — OBJETIVO ──
    secao(doc, '1', 'OBJETIVO')
    add_par(doc,
        f'Esta Folha de Extensão de Aplicabilidade formaliza a validade do {codigo} ({revisao}) — '
        f'{titulo} — para o Bloco D do contrato EDISER/Petrobras '
        f'(ICJ 5900.0132378.25.2), seguindo o mesmo padrão adotado para o Bloco B '
        f'(Área Liberada N°002/2026 — PE-2ACO-00005 / Anexo E).')

    # ── SEÇÃO 2 — DOCUMENTO DE REFERÊNCIA ──
    secao(doc, '2', 'DOCUMENTO DE REFERÊNCIA')
    t2 = doc.add_table(rows=4, cols=2)
    t2.style = 'Table Grid'
    dados_ref = [
        ('Código do Documento',  codigo),
        ('Revisão em Vigor',     revisao),
        ('Título',               titulo),
        ('Bloco de Origem',      'Bloco B — Área Liberada N°002/2026'),
    ]
    for i, (k, v) in enumerate(dados_ref):
        bg = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(t2.cell(i,0), bg)
        set_cell_bg(t2.cell(i,1), bg)
        cell_text(t2.cell(i,0), k, bold=True, size=9)
        cell_text(t2.cell(i,1), v, size=9)
    t2.columns[0].width = Cm(6)
    t2.columns[1].width = Cm(12)
    doc.add_paragraph()

    # ── SEÇÃO 3 — DECLARAÇÃO DE EXTENSÃO ──
    secao(doc, '3', 'DECLARAÇÃO DE EXTENSÃO DE APLICABILIDADE')
    add_par(doc,
        f'Declaro que o {codigo} ({revisao}), elaborado e aprovado para as atividades do Bloco B '
        f'do contrato EDISER/Petrobras, possui plena aplicabilidade técnica ao Bloco D, '
        f'uma vez que:')
    itens = [
        'As atividades, riscos, medidas de controle e responsabilidades descritos no documento de referência são equivalentes nos dois blocos;',
        'As normas técnicas e regulamentadoras aplicáveis (NRs, NBRs e N-Petrobras) não sofreram alteração de escopo entre os blocos;',
        'Os adendos da Norma Petrobras N-2869 (ADN-N2869), emitidos em 08/05/2026 e confirmados pela fiscalização em 09/05/2026, são igualmente incorporados a este documento para o Bloco D;',
        'O corpo técnico responsável (Coordenador SMS e Responsável Técnico) permanece o mesmo.',
    ]
    for item in itens:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(item)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    doc.add_paragraph()

    # ── SEÇÃO 4 — PARTICULARIDADES DO BLOCO D ──
    secao(doc, '4', 'PARTICULARIDADES E OBSERVAÇÕES PARA O BLOCO D')
    t3 = doc.add_table(rows=4, cols=2)
    t3.style = 'Table Grid'
    t3.columns[0].width = Cm(5)
    t3.columns[1].width = Cm(13)
    part_dados = [
        ('Bloco de Aplicação',   'Bloco D — Contrato EDISER / ICJ 5900.0132378.25.2'),
        ('Área Liberada',        'A ser emitida — Área Liberada Bloco D (pendente aprovação Petrobras)'),
        ('Adendos N-2869',       'Incorporados conforme folhas ADN-ENG-XXX-N2869, emitidas em 08/05/2026'),
        ('Observações',          'Preencher caso existam particularidades específicas do Bloco D para esta atividade.'),
    ]
    for i, (k, v) in enumerate(part_dados):
        bg = 'FFF8E1' if i % 2 == 0 else 'FFFDE7'
        set_cell_bg(t3.cell(i,0), bg)
        set_cell_bg(t3.cell(i,1), bg)
        cell_text(t3.cell(i,0), k, bold=True, size=9)
        cell_text(t3.cell(i,1), v, size=9)
    doc.add_paragraph()

    # ── SEÇÃO 5 — VIGÊNCIA ──
    secao(doc, '5', 'VIGÊNCIA')
    add_par(doc,
        f'Esta Folha de Extensão entra em vigor na data de sua emissão ({DATA}) e permanece '
        f'válida enquanto o documento de referência estiver vigente e enquanto durar o '
        f'contrato EDISER (ICJ 5900.0132378.25.2).')
    doc.add_paragraph()

    # ── ASSINATURAS ──
    t4 = doc.add_table(rows=2, cols=3)
    t4.style = 'Table Grid'
    t4.columns[0].width = Cm(6)
    t4.columns[1].width = Cm(6)
    t4.columns[2].width = Cm(6)

    set_cell_bg(t4.cell(0,0), '0F2D52')
    set_cell_bg(t4.cell(0,1), '0F2D52')
    set_cell_bg(t4.cell(0,2), '0F2D52')
    cell_text(t4.cell(0,0), 'Elaborado por', bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t4.cell(0,1), 'Aprovado por', bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t4.cell(0,2), 'Data de Vigência', bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    cell_text(t4.cell(1,0),
              'Fábio Martins Nascimento\nCREA-SE 2718664207\nCoordenador de SMS',
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t4.cell(1,1),
              'Lúcio Wagner Saturno Silva\nCREA-SE 270237084-5\nResponsável Técnico / Preposto',
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t4.cell(1,2),
              DATA,
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── SALVAR ──
    nome_arquivo = f'EXT-{codigo}-BLOCO-D.docx'
    caminho = os.path.join(pasta_saida, nome_arquivo)
    doc.save(caminho)
    return nome_arquivo

# ── EXECUÇÃO ───────────────────────────────────────────────────────────────
print('Gerando Folhas de Extensão — APRs...')
for codigo, revisao, titulo in APRS:
    nome = gerar_folha(codigo, revisao, titulo, 'APR', OUT_APR)
    print(f'  ✅ {nome}')

print(f'\nTotal APRs: {len(APRS)} documentos')

print('\nGerando Folhas de Extensão — PEXes...')
for codigo, revisao, titulo in PEXES:
    nome = gerar_folha(codigo, revisao, titulo, 'PEX', OUT_PEX)
    print(f'  ✅ {nome}')

print(f'\nTotal PEXes: {len(PEXES)} documentos')
print(f'\nTotal geral: {len(APRS) + len(PEXES)} documentos gerados.')
print(f'\nPastas:')
print(f'  APRs  → {OUT_APR}')
print(f'  PEXes → {OUT_PEX}')
