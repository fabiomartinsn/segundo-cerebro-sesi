from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Cores JFX ──────────────────────────────────────────────
NAVY   = RGBColor(0x0F, 0x2D, 0x52)
ORANGE = RGBColor(0xE0, 0x7B, 0x2A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF4, 0xF7, 0xFB)   # fundo linhas ímpares
DARK_ORANGE = RGBColor(0xC0, 0x65, 0x18)

FONT = "Calibri"

# ── Helpers ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=10, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.color.rgb = color or RGBColor(0x1A, 0x1A, 0x1A)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = NAVY if level == 1 else ORANGE
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.4)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    return p

def divider(doc, color_hex="0F2D52"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── Documento ──────────────────────────────────────────────
doc = Document()

# Margens
for sec in doc.sections:
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── Cabeçalho ──────────────────────────────────────────────
hdr_tbl = doc.add_table(rows=1, cols=2)
hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_tbl.allow_autofit = False

col_widths = [Cm(11), Cm(6)]
for i, w in enumerate(col_widths):
    hdr_tbl.columns[i].width = w

left  = hdr_tbl.cell(0, 0)
right = hdr_tbl.cell(0, 1)

set_cell_bg(left,  "0F2D52")
set_cell_bg(right, "E07B2A")

left.paragraphs[0].clear()
p_l = left.paragraphs[0]
p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = p_l.add_run("JFX ENGENHARIA  |  GESTÃO SMS\n")
r1.bold = True; r1.font.name = FONT; r1.font.size = Pt(13); r1.font.color.rgb = WHITE
r2 = p_l.add_run("EDISER / Petrobras — Blocos B, D e Q")
r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xCC, 0xD9, 0xEA)
left._tc.get_or_add_tcPr()

right.paragraphs[0].clear()
p_r = right.paragraphs[0]
p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_r.add_run("PLANO DE AÇÃO\n")
r3.bold = True; r3.font.name = FONT; r3.font.size = Pt(12); r3.font.color.rgb = WHITE
r4 = p_r.add_run("04 de maio de 2026")
r4.font.name = FONT; r4.font.size = Pt(10); r4.font.color.rgb = WHITE

doc.add_paragraph()

# ── Título principal ───────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(4)
t.paragraph_format.space_after  = Pt(2)
run_t = t.add_run("PRIORIDADES 04/05/2026 — REGULARIZAÇÃO DA GESTÃO SMS")
run_t.bold = True; run_t.font.name = FONT; run_t.font.size = Pt(14)
run_t.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
run_s = sub.add_run("Plano de retomada pós-reunião 30/04/2026  ·  Origem: cobrança Petrobras/EDISER")
run_s.italic = True; run_s.font.name = FONT; run_s.font.size = Pt(9)
run_s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

divider(doc)

# ── Seção 1 — Contexto ─────────────────────────────────────
add_heading(doc, "1. CONTEXTO E MOTIVAÇÃO")
add_body(doc,
    "Na reunião de Gestão SMS realizada em 30/04/2026, foram identificadas lacunas críticas no sistema "
    "de gestão: ausência de indicadores formalizados, programas não iniciados (VCP, Amigos do Peito, "
    "Projeto EPI) e necessidade de apresentar números na próxima reunião. Este documento estrutura as "
    "ações para regularização completa em três fases até 23/05/2026."
)

# ── Seção 2 — Prioridades do dia ──────────────────────────
add_heading(doc, "2. PRIORIDADES PARA 04/05/2026")
divider(doc, "E07B2A")

# Tabela de prioridades
pri_data = [
    ("P1", "CRÍTICO", "Levantar HHT e dias sem acidente",
     "Contatar Roberta Muniz (TST Adm). Consolidar HHT desde início do contrato. "
     "Confirmar data do último acidente ou ausência. "
     "Ação desbloqueadora: libera cálculo de TFCA, TAR, TOA e entrega do REM abr/2026."),
    ("P2", "URGENTE", "Entregar REM de abril/2026",
     "Prazo: até 05/05/2026 (amanhã). Roberta Muniz responsável. "
     "Verificar DDS de 18/04 a 03/05 com Roberta — lançamento retroativo se necessário. "
     "Calcular indicadores sobre HHT levantado no P1."),
    ("P3", "URGENTE", "Retomar DDS documentado",
     "Reiniciar registro diário no Checkbits a partir de 05/05/2026. "
     "Garantir formulário com: data, tema, participantes e assinaturas. "
     "34 evidências já registradas (02/03–17/04/2026) — base sólida para auditoria."),
    ("P4", "IMPORTANTE", "Iniciar o VCP (Verificação Comportamental Preventiva)",
     "Prazo interno: até 09/05/2026. Definir: formulário, responsável (Sheila Ferreira — TST Campo), "
     "frequência e meta. VCP é indicador proativo mensal cobrado pela Petrobras."),
]

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Larguras
widths = [Cm(1.2), Cm(2.2), Cm(4.5), Cm(9.0)]
for i, w in enumerate(widths):
    for row in tbl.rows:
        row.cells[i].width = w

# Cabeçalho
headers = ["#", "STATUS", "AÇÃO", "DESCRIÇÃO / RESPONSÁVEL"]
for i, h in enumerate(headers):
    cell_text(tbl.cell(0, i), h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl.cell(0, i), "0F2D52")

for idx, (num, status, acao, desc) in enumerate(pri_data):
    row = tbl.add_row()
    bg = "F4F7FB" if idx % 2 == 0 else "FFFFFF"
    status_color = "C80000" if status == "CRÍTICO" else ("E07B2A" if status == "URGENTE" else "2E6DA4")
    cell_text(row.cells[0], num,    bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[1], status, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[1], status_color.replace("#",""))
    cell_text(row.cells[2], acao,   bold=True, size=10)
    cell_text(row.cells[3], desc,   size=9)
    for c in [row.cells[0], row.cells[2], row.cells[3]]:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ── Seção 3 — Plano de Ação por Fases ─────────────────────
add_heading(doc, "3. PLANO DE AÇÃO — REGULARIZAÇÃO SMS")
add_body(doc, "Estruturado em três fases de acordo com criticidade e dependências técnicas:")
divider(doc, "E07B2A")

# ── Fase 1 ─────────────────────────────────────────────────
add_heading(doc, "FASE 1 — Base de Dados e Indicadores  [04 a 09/05/2026]", level=2)
add_body(doc, "Desbloqueador de todos os demais indicadores — Ciclo CHECK (ISO 45001 §9.1).", indent=True)

fase1_data = [
    ("1.1", "04–05/05", "HHT e dias sem acidente",
     "Roberta Muniz", "Aberto",
     "Consolidar HHT desde o início do contrato. Confirmar data do último acidente. "
     "Destrava: TFCA, TAR, TOA e REM."),
    ("1.2", "05/05",    "REM de abril/2026",
     "Roberta Muniz", "PRAZO HOJE",
     "Calcular os 15 indicadores sobre HHT. Entregar à Petrobras até 05/05. "
     "Verificar e lançar DDS retroativo 18/04–03/05 se necessário."),
    ("1.3", "05/05→",   "Retomada DDS diário",
     "Sheila Ferreira / TST", "Aberto",
     "Registro diário no Checkbits. Formulário com data, tema, participantes e assinaturas."),
    ("1.4", "até 09/05","VCP — implantação",
     "Sheila Ferreira", "Aberto",
     "Definir formulário VCP, responsável, frequência mensal e meta de % comportamentos seguros."),
]

tbl2 = doc.add_table(rows=1, cols=6)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
widths2 = [Cm(0.8), Cm(1.8), Cm(3.2), Cm(2.5), Cm(1.8), Cm(6.8)]
for i, w in enumerate(widths2):
    for row in tbl2.rows:
        row.cells[i].width = w

hdrs2 = ["Item", "Prazo", "Ação", "Responsável", "Status", "Detalhamento"]
for i, h in enumerate(hdrs2):
    cell_text(tbl2.cell(0, i), h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl2.cell(0, i), "0F2D52")

for idx, (item, prazo, acao, resp, status, det) in enumerate(fase1_data):
    row = tbl2.add_row()
    bg = "FFF5EC" if idx % 2 == 0 else "FFFFFF"
    st_color = "C80000" if "HOJE" in status else "2E6DA4"
    cell_text(row.cells[0], item,   bold=True, size=9,  align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[1], prazo,  bold=True, color=RGBColor(0xE0,0x7B,0x2A), size=9)
    cell_text(row.cells[2], acao,   bold=True, size=9)
    cell_text(row.cells[3], resp,   size=9)
    cell_text(row.cells[4], status, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[4], st_color)
    cell_text(row.cells[5], det,    size=9)
    for c in [row.cells[0], row.cells[1], row.cells[2], row.cells[3], row.cells[5]]:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ── Fase 2 ─────────────────────────────────────────────────
add_heading(doc, "FASE 2 — Conformidade Legal e Operacional  [12 a 16/05/2026]", level=2)
add_body(doc, "Redução de passivo legal e risco de auditoria Petrobras — Ciclo DO (§2.2, §2.5).", indent=True)

fase2_data = [
    ("2.1", "12–14/05", "Matriz de certificados NRs",
     "Roberta Muniz", "Aberto",
     "Mapear todos os certificados (NR-10 SEP, NR-33, NR-35, NR-12). "
     "Identificar vencidos e a vencer em 30 dias. Gerar matriz com validades."),
    ("2.2", "15–16/05", "Revisão APRs atividades críticas",
     "Sheila Ferreira", "Aberto",
     "Verificar APR aprovada para cada atividade crítica em execução: "
     "Trabalho em Altura (NR-35), Espaço Confinado (NR-33), Elétrica (NR-10). "
     "Atualizar documentos desatualizados."),
]

tbl3 = doc.add_table(rows=1, cols=6)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate(widths2):
    for row in tbl3.rows:
        row.cells[i].width = w

for i, h in enumerate(hdrs2):
    cell_text(tbl3.cell(0, i), h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl3.cell(0, i), "0F2D52")

for idx, (item, prazo, acao, resp, status, det) in enumerate(fase2_data):
    row = tbl3.add_row()
    bg = "F0F5FF" if idx % 2 == 0 else "FFFFFF"
    cell_text(row.cells[0], item,   bold=True, size=9,  align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[1], prazo,  bold=True, color=RGBColor(0xE0,0x7B,0x2A), size=9)
    cell_text(row.cells[2], acao,   bold=True, size=9)
    cell_text(row.cells[3], resp,   size=9)
    cell_text(row.cells[4], status, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[4], "2E6DA4")
    cell_text(row.cells[5], det,    size=9)
    for c in [row.cells[0], row.cells[1], row.cells[2], row.cells[3], row.cells[5]]:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ── Fase 3 ─────────────────────────────────────────────────
add_heading(doc, "FASE 3 — Programas Cobrados na Reunião  [19 a 23/05/2026]", level=2)
add_body(doc, "Implantação dos programas apontados como não iniciados — Ciclo PLAN/DO.", indent=True)

fase3_data = [
    ("3.1", "19–21/05", "Amigos do Peito",
     "Coord. SMS + Sheila", "Aberto",
     "Planejar cronograma, designar responsável, definir público-alvo e periodicidade semanal."),
    ("3.2", "22–23/05", "Projeto EPI",
     "Roberta Muniz + Sheila", "Aberto",
     "Inventário atual de EPIs, estado de conservação, registros de entrega assinados. "
     "Base: NR-6 + procedimento interno JFX."),
]

tbl4 = doc.add_table(rows=1, cols=6)
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate(widths2):
    for row in tbl4.rows:
        row.cells[i].width = w

for i, h in enumerate(hdrs2):
    cell_text(tbl4.cell(0, i), h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl4.cell(0, i), "0F2D52")

for idx, (item, prazo, acao, resp, status, det) in enumerate(fase3_data):
    row = tbl4.add_row()
    bg = "F5FFF5" if idx % 2 == 0 else "FFFFFF"
    cell_text(row.cells[0], item,   bold=True, size=9,  align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[1], prazo,  bold=True, color=RGBColor(0xE0,0x7B,0x2A), size=9)
    cell_text(row.cells[2], acao,   bold=True, size=9)
    cell_text(row.cells[3], resp,   size=9)
    cell_text(row.cells[4], status, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[4], "2E6DA4")
    cell_text(row.cells[5], det,    size=9)
    for c in [row.cells[0], row.cells[1], row.cells[2], row.cells[3], row.cells[5]]:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ── Seção 4 — Entregáveis ──────────────────────────────────
add_heading(doc, "4. ENTREGÁVEIS PARA A PRÓXIMA REUNIÃO PETROBRAS")
divider(doc, "E07B2A")

ent_data = [
    ("HHT acumulado (desde início do contrato)",          "Fase 1 — item 1.1"),
    ("Dias sem acidente",                                  "Fase 1 — item 1.1"),
    ("Indicadores reativos: TFCA · TAR · TOA",            "Calculados sobre HHT"),
    ("DDS realizados no mês (evidências Checkbits)",       "Fase 1 — item 1.3"),
    ("REM de abril/2026 entregue",                        "Fase 1 — item 1.2"),
    ("VCP iniciado + primeiros dados observados",          "Fase 1 — item 1.4"),
    ("Matriz de certificados com validades",               "Fase 2 — item 2.1"),
    ("Cronograma fases 2 e 3 com responsáveis e prazos",  "Este documento"),
]

tbl5 = doc.add_table(rows=1, cols=2)
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl5.columns[0].width = Cm(12)
tbl5.columns[1].width = Cm(5)

cell_text(tbl5.cell(0,0), "ENTREGÁVEL", bold=True, color=WHITE, size=9)
cell_text(tbl5.cell(0,1), "ORIGEM", bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(tbl5.cell(0,0), "0F2D52")
set_cell_bg(tbl5.cell(0,1), "0F2D52")

for idx, (entrega, origem) in enumerate(ent_data):
    row = tbl5.add_row()
    bg = "F4F7FB" if idx % 2 == 0 else "FFFFFF"
    cell_text(row.cells[0], entrega, size=10)
    cell_text(row.cells[1], origem, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)

doc.add_paragraph()

# ── Seção 5 — Equipe SMS ───────────────────────────────────
add_heading(doc, "5. EQUIPE SMS — CONTRATO EDISER")
divider(doc, "E07B2A")

equipe_data = [
    ("Fábio Martins Nascimento", "Coordenador SMS", "Gestão geral, interface Petrobras, liderança do plano"),
    ("Roberta Muniz",            "TST Administrativa", "REM, HHER, indicadores, treinamentos, CAT, CIPA"),
    ("Sheila Ferreira",          "TST Campo", "DDS, APRs, VCP, AUDICOMP, inspeções de campo"),
    ("Lúcio Saturno",            "Preposto JFX", "Liderança, recursos, reuniões formais com Petrobras"),
]

tbl6 = doc.add_table(rows=1, cols=3)
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl6.columns[0].width = Cm(5)
tbl6.columns[1].width = Cm(4)
tbl6.columns[2].width = Cm(8)

hdrs6 = ["NOME", "FUNÇÃO", "ATRIBUIÇÃO PRINCIPAL"]
for i, h in enumerate(hdrs6):
    cell_text(tbl6.cell(0,i), h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl6.cell(0,i), "0F2D52")

for idx, (nome, func, atrib) in enumerate(equipe_data):
    row = tbl6.add_row()
    bg = "F4F7FB" if idx % 2 == 0 else "FFFFFF"
    cell_text(row.cells[0], nome,  bold=True, size=9)
    cell_text(row.cells[1], func,  size=9)
    cell_text(row.cells[2], atrib, size=9)
    for c in row.cells:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ── Rodapé ─────────────────────────────────────────────────
divider(doc)
rodape = doc.add_paragraph()
rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
rodape.paragraph_format.space_before = Pt(4)
r_run = rodape.add_run(
    "JFX Engenharia — Coordenação de SMS  |  ICJ 5900.0132378.25.2  |  EDISER / Petrobras\n"
    "Documento gerado em 03/05/2026  ·  Uso interno"
)
r_run.font.name = FONT
r_run.font.size = Pt(8)
r_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# ── Salvar ─────────────────────────────────────────────────
out = r"C:\Users\fabio\Documents\segundo-cerebro\Plano-Acao-SMS-04052026.docx"
doc.save(out)
print(f"OK: {out}")
